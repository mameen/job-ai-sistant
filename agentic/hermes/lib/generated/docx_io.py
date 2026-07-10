"""DOCX create/edit for Application Coach outputs.

Libraries (free, pip-installable):
- **python-docx** (MIT) — create documents, edit paragraphs/tables/styles
- **docxtpl** (LGPL-2.1) — Jinja2 merge into existing .docx templates (keeps layout)

Templates (Project Career Zazu):
  .kb/templates/resume/pm-resume.docx
  .kb/templates/cover/cover.docx
  .kb/templates/brief/application-brief.docx
  manifest: .kb/templates/manifest.yaml
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Any

# Legacy fallbacks when .kb/templates/ not yet built.
RESUME_TEMPLATE_REL = "private/originals/resume-repo/pm-resume.docx"
COVER_TEMPLATE_REL = "private/originals/resume-repo/cover.docx"

ZAZU_RESUME_TEMPLATE_REL = "templates/resume/pm-resume.docx"
ZAZU_COVER_TEMPLATE_REL = "templates/cover/cover.docx"
ZAZU_BRIEF_TEMPLATE_REL = "templates/brief/application-brief.docx"

_STYLE_BODY = "Body Text"
_STYLE_NORMAL = "Normal"
_STYLE_H1 = "Heading 1"
_STYLE_H2 = "Heading 2"
_STYLE_H3 = "Heading 3"
_STYLE_BULLET = "List Bullet"

_SECTION_MARKER_RE = re.compile(r"^\{\{SECTION:([A-Z0-9_]+)\}\}$")

_COVER_BRACKET_TOKENS = (
    "[COMPANY]",
    "[ROLE TITLE]",
    "[JOB TITLE]",
    "[TITLE]",
)

_BRIEF_BRACKET_TOKENS = (
    "[COMPANY]",
    "[ROLE TITLE]",
    "[JOB_ID]",
    "[JOB_DATE]",
    "[URL]",
)

# Brief template Heading 2 labels → patch section ids (heading-anchored merge fallback).
BRIEF_HEADING_FOR_SECTION: dict[str, str] = {
    "Opportunity Intelligence": "OPPORTUNITY_INTELLIGENCE",
    "Flag Analysis": "FLAG_ANALYSIS",
    "Level Mapping": "LEVEL_MAPPING",
    "Employer Summary": "EMPLOYER_SUMMARY",
    "Top Matching Experience": "TOP_MATCHING_EXPERIENCE",
    "JD Alignment": "JD_ALIGNMENT",
    "Gaps and Mitigations": "GAPS_MITIGATIONS",
    "STAR Stories to Prepare": "STAR_STORIES",
    "Interview Intelligence": "INTERVIEW_INTELLIGENCE",
    "Compensation and Next Steps": "COMPENSATION_NEXT_STEPS",
    "Recommendation": "RECOMMENDATION",
}

# Optional resume sections — omitted patch content removes header + placeholder.
RESUME_OPTIONAL_SECTIONS = frozenset(
    {"SELECTED_TECHNICAL", "CERTIFICATIONS", "SELECTED_THEMES", "TECH_STACK"}
)


def section_marker(section_id: str) -> str:
    return f"{{{{SECTION:{section_id}}}}}"


def substitute_bracket_placeholders(
    text: str,
    *,
    company: str,
    job_title: str,
    job_id: str = "",
    job_date: str = "",
    url: str = "",
) -> str:
    """Replace bracket tokens in cover/brief text."""
    out = text
    out = out.replace("[COMPANY]", company)
    out = out.replace("[ROLE TITLE]", job_title)
    out = out.replace("[JOB TITLE]", job_title)
    out = out.replace("[TITLE]", job_title)
    out = out.replace("[JOB_ID]", job_id or "na")
    out = out.replace("[JOB_DATE]", job_date)
    out = out.replace("[URL]", url or "na")
    return out


def parse_section_markdown(markdown: str) -> dict[str, str]:
    """Parse ``## SECTION_ID`` blocks into a section-id → body map."""
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for raw in markdown.splitlines():
        line = raw.rstrip()
        heading = re.match(r"^##\s+([A-Z0-9_]+)\s*$", line.strip())
        if heading:
            current = heading.group(1)
            sections.setdefault(current, [])
            continue
        if current is None:
            continue
        sections[current].append(line)
    return {
        key: "\n".join(lines).strip()
        for key, lines in sections.items()
        if "\n".join(lines).strip()
    }


HERMES_PKG = Path(__file__).resolve().parents[2]
SCAFFOLD_TEMPLATES = HERMES_PKG / "kb" / "scaffold" / "templates"


def resolve_zazu_template(kb_root: Path, kind: str) -> Path:
    """Return Zazu template path (vault first, then tracked scaffold, then legacy)."""
    candidates: list[Path] = []
    rel = {
        "resume": (ZAZU_RESUME_TEMPLATE_REL, "resume/pm-resume.docx"),
        "cover": (ZAZU_COVER_TEMPLATE_REL, "cover/cover.docx"),
        "brief": (ZAZU_BRIEF_TEMPLATE_REL, "brief/application-brief.docx"),
    }.get(kind)
    if rel:
        kb_rel, scaffold_rel = rel
        candidates.append(kb_root / kb_rel)
        candidates.append(SCAFFOLD_TEMPLATES / scaffold_rel)
    legacy = {
        "resume": RESUME_TEMPLATE_REL,
        "cover": COVER_TEMPLATE_REL,
    }.get(kind)
    if legacy:
        candidates.append(kb_root / legacy)
    for path in candidates:
        if path.is_file():
            return path
    raise FileNotFoundError(f"DOCX template missing for kind={kind!r} under {kb_root}")


def fill_bracket_tokens_in_docx(
    docx_path: Path,
    *,
    company: str,
    job_title: str,
    job_id: str = "",
    job_date: str = "",
    url: str = "",
    tokens: tuple[str, ...] = _COVER_BRACKET_TOKENS,
) -> Path:
    """Replace bracket placeholders in an on-disk DOCX (preserves paragraph styles)."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError("pip install python-docx") from exc

    doc = Document(str(docx_path))
    changed = False

    def _rewrite_paragraph(para) -> None:
        nonlocal changed
        original = para.text
        updated = substitute_bracket_placeholders(
            original,
            company=company,
            job_title=job_title,
            job_id=job_id,
            job_date=job_date,
            url=url,
        )
        if updated == original:
            return
        changed = True
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = updated
        else:
            para.add_run(updated)

    for para in doc.paragraphs:
        if any(token in para.text for token in tokens):
            _rewrite_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if any(token in para.text for token in tokens):
                        _rewrite_paragraph(para)

    if changed:
        doc.save(str(docx_path))
    return docx_path


def fill_cover_placeholders(
    docx_path: Path,
    *,
    company: str,
    job_title: str,
) -> Path:
    """Replace cover bracket placeholders in an on-disk cover letter."""
    return fill_bracket_tokens_in_docx(
        docx_path,
        company=company,
        job_title=job_title,
        tokens=_COVER_BRACKET_TOKENS,
    )


def _remove_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def _insert_paragraph_after(paragraph, text: str, style) -> Any:
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    from docx.text.paragraph import Paragraph  # type: ignore[import-untyped]

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def _parsed_section_lines(body: str) -> list[tuple[str, str]]:
    """Parse patch body lines; ``###`` → Heading 3 subhead, ``-`` → bullet."""
    items: list[tuple[str, str]] = []
    for raw in body.splitlines():
        stripped = raw.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            items.append(("h3", stripped[4:].strip()))
        elif stripped.startswith("- "):
            items.append(("bullet", stripped[2:].strip()))
        else:
            items.append(("normal", stripped))
    return items


def _style_for_line_kind(doc, base_style, kind: str):
    if kind == "h3":
        try:
            return doc.styles["Heading 3"]
        except KeyError:
            pass
    if kind == "bullet":
        try:
            return doc.styles[_STYLE_BULLET]
        except KeyError:
            pass
    return base_style


def _section_lines(body: str) -> list[str]:
    return [text for _, text in _parsed_section_lines(body)]


def _insert_section_body_after(
    doc,
    anchor_para,
    body: str,
    *,
    bracket_ctx: dict[str, str],
    base_style,
) -> None:
    """Insert parsed patch lines after ``anchor_para``."""
    company = bracket_ctx.get("company", "")
    job_title = bracket_ctx.get("job_title", "")
    job_id = bracket_ctx.get("job_id", "")
    job_date = bracket_ctx.get("job_date", "")
    url = bracket_ctx.get("url", "")

    parsed = _parsed_section_lines(body)
    if not parsed:
        return

    first_kind, first_text = parsed[0]
    style = _style_for_line_kind(doc, base_style, first_kind)
    first = substitute_bracket_placeholders(
        first_text,
        company=company,
        job_title=job_title,
        job_id=job_id,
        job_date=job_date,
        url=url,
    )
    anchor = _insert_paragraph_after(anchor_para, first, style)
    for kind, line in parsed[1:]:
        rendered = substitute_bracket_placeholders(
            line,
            company=company,
            job_title=job_title,
            job_id=job_id,
            job_date=job_date,
            url=url,
        )
        line_style = _style_for_line_kind(doc, base_style, kind)
        anchor = _insert_paragraph_after(anchor, rendered, line_style)


def _merge_brief_patch_by_headings(
    doc,
    sections: dict[str, str],
    *,
    bracket_ctx: dict[str, str],
) -> int:
    """Replace brief body under each Heading 2 when ``{{SECTION:}}`` markers are gone."""
    merged = 0
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        heading = para.text.strip()
        section_id = BRIEF_HEADING_FOR_SECTION.get(heading)
        if not section_id:
            i += 1
            continue
        body = sections.get(section_id, "").strip()
        if not body:
            i += 1
            continue

        j = i + 1
        while j < len(doc.paragraphs):
            next_para = doc.paragraphs[j]
            next_text = next_para.text.strip()
            if next_para.style.name in ("Heading 1", "Heading 2") and next_text:
                break
            j += 1

        for k in range(j - 1, i, -1):
            _remove_paragraph(doc.paragraphs[k])

        base_style = doc.styles[_STYLE_NORMAL]
        _insert_section_body_after(
            doc, para, body, bracket_ctx=bracket_ctx, base_style=base_style
        )
        merged += 1
        i += 1

    return merged


def merge_section_patch_into_docx(
    docx_path: Path,
    sections: dict[str, str],
    *,
    optional_sections: frozenset[str] = frozenset(),
    bracket_ctx: dict[str, str] | None = None,
    brief_heading_fallback: bool = False,
) -> int:
    """Replace ``{{SECTION:ID}}`` markers with patch content; preserve template styles.

    Returns the number of sections merged. When ``brief_heading_fallback`` is True
    and no marker merges occurred, retries using Heading 2 labels in the brief template.
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError("pip install python-docx") from exc

    doc = Document(str(docx_path))
    bracket_ctx = bracket_ctx or {}
    company = bracket_ctx.get("company", "")
    job_title = bracket_ctx.get("job_title", "")
    job_id = bracket_ctx.get("job_id", "")
    job_date = bracket_ctx.get("job_date", "")
    url = bracket_ctx.get("url", "")

    merged = 0
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()
        match = _SECTION_MARKER_RE.match(text)
        if not match:
            i += 1
            continue

        section_id = match.group(1)
        body = sections.get(section_id, "").strip()
        if not body and section_id in optional_sections:
            if i > 0:
                _remove_paragraph(doc.paragraphs[i - 1])
                i -= 1
            _remove_paragraph(doc.paragraphs[i])
            continue

        parsed = _parsed_section_lines(body)
        if not parsed:
            i += 1
            continue

        first_kind, first_text = parsed[0]
        style = _style_for_line_kind(doc, para.style, first_kind)
        first = substitute_bracket_placeholders(
            first_text,
            company=company,
            job_title=job_title,
            job_id=job_id,
            job_date=job_date,
            url=url,
        )
        para.clear()
        para.add_run(first)
        para.style = style
        anchor = para
        for kind, line in parsed[1:]:
            rendered = substitute_bracket_placeholders(
                line,
                company=company,
                job_title=job_title,
                job_id=job_id,
                job_date=job_date,
                url=url,
            )
            line_style = _style_for_line_kind(doc, para.style, kind)
            anchor = _insert_paragraph_after(anchor, rendered, line_style)
        merged += 1
        i += 1

    if brief_heading_fallback and merged == 0 and sections:
        merged = _merge_brief_patch_by_headings(doc, sections, bracket_ctx=bracket_ctx)

    doc.save(str(docx_path))
    fill_bracket_tokens_in_docx(
        docx_path,
        company=company,
        job_title=job_title,
        job_id=job_id,
        job_date=job_date,
        url=url,
        tokens=_COVER_BRACKET_TOKENS + _BRIEF_BRACKET_TOKENS,
    )
    return merged


def coach_sidecar_status(run_dir: Path) -> dict[str, bool]:
    def _has(kind: str) -> bool:
        return any(run_dir.glob(f"*_{kind}_patch.md")) or any(run_dir.glob(f"*_{kind}.md"))

    return {kind: _has(kind) for kind in ("resume", "cover", "brief")}


def materialize_markdown_sidecars(
    run_dir: Path,
    *,
    company: str = "",
    job_title: str = "",
    job_id: str = "",
    job_date: str = "",
    url: str = "",
) -> list[Path]:
    """Merge coach ``*_patch.md`` (preferred) or legacy ``*_{kind}.md`` into DOCX."""
    updated: list[Path] = []
    bracket_ctx = {
        "company": company,
        "job_title": job_title,
        "job_id": job_id,
        "job_date": job_date,
        "url": url,
    }

    for kind, optional in (
        ("resume", RESUME_OPTIONAL_SECTIONS),
        ("cover", frozenset()),
        ("brief", frozenset()),
    ):
        for docx_path in sorted(run_dir.glob(f"*_{kind}.docx")):
            patch_path = docx_path.with_name(docx_path.stem + "_patch.md")
            legacy_path = docx_path.with_suffix(".md")
            md_path = patch_path if patch_path.is_file() else legacy_path
            if not md_path.is_file():
                continue
            body = md_path.read_text(encoding="utf-8").strip()
            if not body:
                continue

            sections = parse_section_markdown(body)
            if sections:
                count = merge_section_patch_into_docx(
                    docx_path,
                    sections,
                    optional_sections=optional,
                    bracket_ctx=bracket_ctx,
                    brief_heading_fallback=(kind == "brief"),
                )
                if count > 0:
                    updated.append(docx_path)
                continue

            if kind == "brief":
                write_markdown_docx(docx_path, body)
                updated.append(docx_path)
            else:
                header_lines = 4 if kind == "resume" else 2
                clone_and_fill_body(
                    docx_path,
                    docx_path,
                    body,
                    header_lines=header_lines,
                )
                updated.append(docx_path)

    return updated


def kb_template(kb_root: Path, rel: str) -> Path:
    path = kb_root / rel
    if not path.is_file():
        raise FileNotFoundError(f"DOCX template missing: {path}")
    return path


def copy_template(template_path: Path, dest_path: Path) -> Path:
    """Copy a .docx template verbatim (preserves all styles/layout)."""
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, dest_path)
    return dest_path


def render_from_template(
    template_path: Path,
    context: dict[str, Any],
    dest_path: Path,
) -> Path:
    """Render Jinja placeholders in a Word template (preserves styles/layout)."""
    try:
        from docxtpl import DocxTemplate  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError("pip install docxtpl") from exc

    dest_path.parent.mkdir(parents=True, exist_ok=True)
    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    tpl.save(str(dest_path))
    return dest_path


def _style_or_normal(doc, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles[_STYLE_NORMAL]


def _add_inline_markdown(paragraph, text: str) -> None:
    """Support ``**bold**`` spans inside a paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def write_markdown_docx(dest_path: Path, markdown: str) -> Path:
    """Write markdown-ish text with Word heading/list styles (legacy brief path)."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError("pip install python-docx") from exc

    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style=_STYLE_H3)
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style=_STYLE_H2)
        elif stripped.startswith("# "):
            doc.add_paragraph(stripped[2:].strip(), style=_STYLE_H1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style=_STYLE_BULLET)
        else:
            para = doc.add_paragraph(style=_STYLE_NORMAL)
            _add_inline_markdown(para, stripped)

    doc.save(str(dest_path))
    return dest_path


def write_plain_docx(
    dest_path: Path,
    *,
    title: str | None = None,
    paragraphs: list[str],
) -> Path:
    """Create a simple DOCX (prefer template merge for briefs)."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError("pip install python-docx") from exc

    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_paragraph(title, style=_STYLE_H1)
    for block in paragraphs:
        for line in block.splitlines():
            stripped = line.strip()
            if stripped:
                doc.add_paragraph(stripped, style=_STYLE_NORMAL)
    doc.save(str(dest_path))
    return dest_path


def clone_and_fill_body(
    template_path: Path,
    dest_path: Path,
    body_text: str,
    *,
    clear_extra_paragraphs: bool = True,
    header_lines: int = 2,
) -> Path:
    """Legacy: copy template and replace all body text (destroys layout)."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError("pip install python-docx") from exc

    dest_path.parent.mkdir(parents=True, exist_ok=True)
    if template_path.resolve() != dest_path.resolve():
        shutil.copy2(template_path, dest_path)
    doc = Document(str(dest_path))

    lines = [ln for ln in body_text.splitlines() if ln.strip()]
    if not lines:
        doc.save(str(dest_path))
        return dest_path

    normal_style = _style_or_normal(doc, _STYLE_NORMAL)
    body_style = _style_or_normal(doc, _STYLE_BODY)

    for para in doc.paragraphs:
        para.clear()

    for i, line in enumerate(lines):
        style = normal_style if i < header_lines else body_style
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            p.add_run(line.strip())
            p.style = style
        else:
            p = doc.add_paragraph(line.strip(), style=style)

    if clear_extra_paragraphs and len(doc.paragraphs) > len(lines):
        for para in list(doc.paragraphs[len(lines) :]):
            el = para._element
            el.getparent().remove(el)

    doc.save(str(dest_path))
    return dest_path


def write_application_triplet(
    *,
    kb_root: Path,
    run_dir: Path,
    company: str,
    job_title: str,
    job_id: str,
    job_date: str,
    resume_text: str,
    cover_text: str,
    brief_text: str,
    use_jinja: bool = False,
    jinja_context: dict[str, Any] | None = None,
    resume_from_template: bool = True,
    cover_from_template: Path | None = None,
) -> list[Path]:
    """Write resume + cover + brief DOCX from Zazu templates into proposals/<run>/."""
    from .naming import artifact_filename

    ctx = jinja_context or {}
    written: list[Path] = []
    bracket_ctx = {
        "company": company,
        "job_title": job_title,
        "job_id": job_id,
        "job_date": job_date,
        "url": ctx.get("url", ""),
    }

    specs: tuple[tuple[str, str], ...] = (
        ("resume", resume_text),
        ("cover", cover_text),
        ("brief", brief_text),
    )

    for kind, text in specs:
        name = artifact_filename(
            company=company,
            job_title=job_title,
            job_id=job_id,
            job_date=job_date,
            kind=kind,
        )
        dest = run_dir / name

        if kind == "brief":
            template = resolve_zazu_template(kb_root, "brief")
            copy_template(template, dest)
            # Keep {{SECTION:}} placeholders for coach merge — only fill header brackets.
            fill_bracket_tokens_in_docx(
                dest,
                company=company,
                job_title=job_title,
                job_id=job_id,
                job_date=job_date,
                url=bracket_ctx["url"],
                tokens=_BRIEF_BRACKET_TOKENS,
            )
        elif use_jinja:
            template = resolve_zazu_template(kb_root, kind)
            render_from_template(template, ctx, dest)
        elif kind == "resume" and resume_from_template:
            copy_template(resolve_zazu_template(kb_root, "resume"), dest)
        elif kind == "cover" and cover_from_template is not None:
            copy_template(cover_from_template, dest)
            fill_cover_placeholders(dest, company=company, job_title=job_title)
        elif kind == "cover" and cover_text.strip() and "{{SECTION:" not in cover_text:
            clone_and_fill_body(
                resolve_zazu_template(kb_root, "cover"),
                dest,
                substitute_bracket_placeholders(
                    cover_text,
                    company=company,
                    job_title=job_title,
                    job_id=job_id,
                    job_date=job_date,
                ),
                header_lines=2,
            )
        elif kind == "cover":
            copy_template(resolve_zazu_template(kb_root, "cover"), dest)
            fill_cover_placeholders(dest, company=company, job_title=job_title)
        else:
            copy_template(resolve_zazu_template(kb_root, kind), dest)

        written.append(dest)
    return written
