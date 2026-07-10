#!/usr/bin/env python3
"""Convert vault resume/cover originals into Project Career Zazu DOCX templates.

Preserves page setup, paragraph styles, and static section headers. Replaces
tailorable body ranges with ``{{SECTION:<id>}}`` markers for deterministic merge.

Run once after updating pm-resume.docx or cover.docx in the vault:

    python agentic/hermes/scripts/build_zazu_templates.py

Outputs:
  agentic/hermes/.kb/templates/resume/pm-resume.docx
  agentic/hermes/.kb/templates/cover/cover.docx
  agentic/hermes/.kb/templates/brief/application-brief.docx
  agentic/hermes/kb/scaffold/templates/...  (tracked copies for bootstrap)
"""

from __future__ import annotations

import shutil
import sys
from pathlib import Path
from typing import Any

REPO = Path(__file__).resolve().parents[3]
HERMES_PKG = REPO / "agentic" / "hermes"
KB_ROOT = HERMES_PKG / ".kb"
SCAFFOLD_TEMPLATES = HERMES_PKG / "kb" / "scaffold" / "templates"
KB_TEMPLATES = KB_ROOT / "templates"

RESUME_SRC = KB_ROOT / "private" / "originals" / "resume-repo" / "pm-resume.docx"
COVER_SRC = KB_ROOT / "private" / "originals" / "resume-repo" / "cover.docx"

_STYLE_BODY = "Body Text"
_STYLE_NORMAL = "Normal"
_PLACEHOLDER_PREFIX = "{{SECTION:"


def _section_marker(section_id: str) -> str:
    return f"{{{{SECTION:{section_id}}}}}"


def _require_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise SystemExit("pip install python-docx") from exc


def _remove_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def _collapse_range(doc, start: int, end: int, placeholder: str, *, style_name: str) -> None:
    """Replace paragraphs [start, end] inclusive with one placeholder paragraph."""
    paras = doc.paragraphs
    if start > end or start >= len(paras):
        return
    anchor = paras[start]
    anchor.text = placeholder
    anchor.style = doc.styles[style_name] if style_name in [s.name for s in doc.styles] else doc.styles["Normal"]
    for para in list(paras[start + 1 : end + 1]):
        _remove_paragraph(para)


def build_resume_template(src: Path, dest: Path) -> None:
    from docx import Document

    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dest)
    doc = Document(str(dest))

    # Indices from pm-resume.docx structure (verified 2026-07-08).
    sections: list[tuple[int, int, str, str]] = [
        (4, 4, "SUMMARY", _STYLE_BODY),
        (7, 12, "AREAS_OF_EXPERTISE", _STYLE_BODY),
        (14, 48, "EXPERIENCES", _STYLE_BODY),
        (50, 55, "SELECTED_TECHNICAL", _STYLE_BODY),
        (58, 60, "EDUCATION", _STYLE_BODY),
        (62, 71, "CERTIFICATIONS", _STYLE_BODY),
        (73, 78, "SELECTED_THEMES", _STYLE_BODY),
        (80, 83, "TECH_STACK", _STYLE_BODY),
    ]
    # Collapse from bottom up so indices stay valid.
    for start, end, section_id, style in reversed(sections):
        _collapse_range(doc, start, end, _section_marker(section_id), style_name=style)

    doc.save(str(dest))
    print(f"  ✓ resume template → {dest.relative_to(REPO)}")


def build_cover_template(src: Path, dest: Path) -> None:
    from docx import Document

    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dest)
    doc = Document(str(dest))

    # Body narrative paragraphs (indices from cover.docx structure).
    cover_paras = [
        (9, "COVER_P1"),
        (11, "COVER_P2"),
        (13, "COVER_P3"),
        (15, "COVER_P4"),
        (17, "COVER_P5"),
    ]
    for idx, section_id in reversed(cover_paras):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            para.text = _section_marker(section_id)
            para.style = doc.styles[_STYLE_BODY]

    doc.save(str(dest))
    print(f"  ✓ cover template → {dest.relative_to(REPO)}")


def _trim_paragraphs_from(doc, start_idx: int) -> None:
    while len(doc.paragraphs) > start_idx:
        _remove_paragraph(doc.paragraphs[start_idx])


def _ensure_brief_heading_styles(doc) -> dict[str, Any]:
    """Add Heading 1–3 to a cloned resume doc (it only ships Normal/Body Text/Heading)."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt

    normal = doc.styles["Normal"]
    names = ("Heading 1", "Heading 2", "Heading 3")
    existing = {s.name for s in doc.styles}
    specs = {
        "Heading 1": ("Franklin Gothic Medium", Pt(14), True),
        "Heading 2": ("Franklin Gothic Book", Pt(11), True),
        "Heading 3": ("Franklin Gothic Book", Pt(10), True),
    }
    for name, (font_name, size, bold) in specs.items():
        if name not in existing:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
            style.font.name = font_name
            style.font.size = size
            style.font.bold = bold
        else:
            style = doc.styles[name]
            style.font.name = font_name
            style.font.size = size
            style.font.bold = bold
    return {name: doc.styles[name] for name in names}


def _add_styled_paragraph(doc, text: str, style_name: str):
    return doc.add_paragraph(text, style=style_name)


def build_brief_template(style_src: Path, dest: Path) -> None:
    """Build brief from resume DOCX — Normal body, Heading 1/2/3 section hierarchy."""
    from docx import Document

    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(style_src, dest)
    doc = Document(str(dest))
    _ensure_brief_heading_styles(doc)
    _trim_paragraphs_from(doc, 2)  # keep name + contact header (Normal, Franklin Gothic)

    _add_styled_paragraph(doc, "", "Normal")
    _add_styled_paragraph(
        doc,
        "Application Brief — [COMPANY] — [ROLE TITLE]",
        "Heading 1",
    )
    _add_styled_paragraph(
        doc,
        "Job ID: [JOB_ID] | Date: [JOB_DATE] | Posting: [URL]",
        "Normal",
    )
    _add_styled_paragraph(doc, "", "Normal")

    brief_sections: list[tuple[str, str]] = [
        ("OPPORTUNITY_INTELLIGENCE", "Opportunity Intelligence"),
        ("FLAG_ANALYSIS", "Flag Analysis"),
        ("LEVEL_MAPPING", "Level Mapping"),
        ("EMPLOYER_SUMMARY", "Employer Summary"),
        ("TOP_MATCHING_EXPERIENCE", "Top Matching Experience"),
        ("JD_ALIGNMENT", "JD Alignment"),
        ("GAPS_MITIGATIONS", "Gaps and Mitigations"),
        ("STAR_STORIES", "STAR Stories to Prepare"),
        ("INTERVIEW_INTELLIGENCE", "Interview Intelligence"),
        ("COMPENSATION_NEXT_STEPS", "Compensation and Next Steps"),
        ("RECOMMENDATION", "Recommendation"),
    ]

    for section_id, heading in brief_sections:
        _add_styled_paragraph(doc, heading, "Heading 2")
        _add_styled_paragraph(doc, _section_marker(section_id), "Normal")
        _add_styled_paragraph(doc, "", "Normal")

    doc.save(str(dest))
    print(f"  ✓ brief template → {dest.relative_to(REPO)}")


def _sync_scaffold(kb_templates: Path) -> None:
    if not kb_templates.is_dir():
        return
    if SCAFFOLD_TEMPLATES.is_dir():
        shutil.rmtree(SCAFFOLD_TEMPLATES)
    shutil.copytree(kb_templates, SCAFFOLD_TEMPLATES)
    print(f"  ✓ scaffold copy → {SCAFFOLD_TEMPLATES.relative_to(REPO)}")


def main() -> int:
    _require_docx()
    print("== build_zazu_templates ==")

    if not RESUME_SRC.is_file():
        print(f"  ERROR missing resume source: {RESUME_SRC}", file=sys.stderr)
        return 1
    if not COVER_SRC.is_file():
        print(f"  ERROR missing cover source: {COVER_SRC}", file=sys.stderr)
        return 1

    KB_TEMPLATES.mkdir(parents=True, exist_ok=True)
    build_resume_template(RESUME_SRC, KB_TEMPLATES / "resume" / "pm-resume.docx")
    build_cover_template(COVER_SRC, KB_TEMPLATES / "cover" / "cover.docx")
    build_brief_template(RESUME_SRC, KB_TEMPLATES / "brief" / "application-brief.docx")

    manifest_src = HERMES_PKG / "kb" / "scaffold" / "templates" / "manifest.yaml"
    if manifest_src.is_file():
        shutil.copy2(manifest_src, KB_TEMPLATES / "manifest.yaml")

    _sync_scaffold(KB_TEMPLATES)
    print("Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
