from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
HERMES_PKG = REPO / "agentic" / "hermes"
if str(HERMES_PKG) not in sys.path:
    sys.path.insert(0, str(HERMES_PKG))

from lib.generated.docx_io import (  # noqa: E402
    fill_cover_placeholders,
    substitute_bracket_placeholders,
    write_markdown_docx,
    write_plain_docx,
)


class DocxIoTests(unittest.TestCase):
    def test_write_plain_docx(self) -> None:
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed")

        with tempfile.TemporaryDirectory() as tmp:
            dest = Path(tmp) / "test_brief.docx"
            write_plain_docx(dest, title="Test Brief", paragraphs=["Line one.", "Line two."])
            self.assertTrue(dest.is_file())
            from docx import Document

            doc = Document(str(dest))
            self.assertTrue(any("Line one" in p.text for p in doc.paragraphs))

    def test_write_markdown_docx_styles(self) -> None:
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed")

        md = """# Application Brief — Acme

## Strategy

- First bullet point.
- Second bullet with **bold** word.

Normal body paragraph.
"""
        with tempfile.TemporaryDirectory() as tmp:
            dest = Path(tmp) / "brief.docx"
            write_markdown_docx(dest, md)
            doc = Document(str(dest))
            styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
            self.assertIn("Heading 1", styles)
            self.assertIn("Heading 2", styles)
            self.assertIn("List Bullet", styles)
            self.assertIn("Normal", styles)
            total = sum(len(p.text) for p in doc.paragraphs)
            self.assertGreater(total, 80)

    def test_fill_cover_placeholders(self) -> None:
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed")

        text = substitute_bracket_placeholders(
            "Role at [COMPANY] as [ROLE TITLE].",
            company="Cisco Systems",
            job_title="Senior Engineering Manager, AI",
        )
        self.assertIn("Cisco Systems", text)
        self.assertNotIn("[COMPANY]", text)

        with tempfile.TemporaryDirectory() as tmp:
            dest = Path(tmp) / "cover.docx"
            write_plain_docx(
                dest,
                paragraphs=[
                    "Header line",
                    "I am applying for the [ROLE TITLE] role at [COMPANY].",
                ],
            )
            fill_cover_placeholders(
                dest,
                company="Cisco Systems",
                job_title="Senior Engineering Manager, AI",
            )
            doc = Document(str(dest))
            body = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Cisco Systems", body)
            self.assertIn("Senior Engineering Manager, AI", body)
            self.assertNotIn("[COMPANY]", body)
            self.assertNotIn("[ROLE TITLE]", body)


if __name__ == "__main__":
    unittest.main()
