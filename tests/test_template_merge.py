from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
HERMES_PKG = REPO / "agentic" / "hermes"
if str(HERMES_PKG) not in sys.path:
    sys.path.insert(0, str(HERMES_PKG))

from lib.generated.docx_io import (  # noqa: E402
    coach_sidecar_status,
    copy_template,
    materialize_markdown_sidecars,
    merge_section_patch_into_docx,
    parse_section_markdown,
    section_marker,
    write_plain_docx,
)

SCAFFOLD_BRIEF = HERMES_PKG / "kb" / "scaffold" / "templates" / "brief" / "application-brief.docx"


class TemplateMergeTests(unittest.TestCase):
    def test_parse_section_markdown(self) -> None:
        md = """## SUMMARY
Tailored summary line.

## EXPERIENCES
Adobe role line one.
Adobe bullet two.
"""
        sections = parse_section_markdown(md)
        self.assertEqual(sections["SUMMARY"], "Tailored summary line.")
        self.assertIn("Adobe role", sections["EXPERIENCES"])

    def test_merge_brief_sections_preserves_headings(self) -> None:
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed")
        if not SCAFFOLD_BRIEF.is_file():
            self.skipTest("scaffold brief template missing — run build_zazu_templates.py")

        patch = """## OPPORTUNITY_INTELLIGENCE
Role: Test Engineer
Company: Acme Corp

## RECOMMENDATION
CONSIDER — strong platform fit; verify comp band.
"""
        with tempfile.TemporaryDirectory() as tmp:
            dest = Path(tmp) / "brief.docx"
            copy_template(SCAFFOLD_BRIEF, dest)
            merge_section_patch_into_docx(
                dest,
                parse_section_markdown(patch),
                bracket_ctx={
                    "company": "Acme Corp",
                    "job_title": "Test Engineer",
                    "job_id": "123",
                    "job_date": "20260708",
                    "url": "https://example.com",
                },
            )
            doc = Document(str(dest))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            self.assertIn("Opportunity Intelligence", text)
            self.assertIn("Acme Corp", text)
            self.assertIn("CONSIDER", text)
            self.assertNotIn(section_marker("OPPORTUNITY_INTELLIGENCE"), text)

    def test_materialize_patch_sidecars(self) -> None:
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed")
        if not SCAFFOLD_BRIEF.is_file():
            self.skipTest("scaffold brief template missing")

        with tempfile.TemporaryDirectory() as tmp:
            run_dir = Path(tmp)
            brief_docx = run_dir / "acme_role_na_20260708_brief.docx"
            copy_template(SCAFFOLD_BRIEF, brief_docx)
            (run_dir / "acme_role_na_20260708_brief_patch.md").write_text(
                "## RECOMMENDATION\nAPPLY — KB-backed AI platform leadership.\n",
                encoding="utf-8",
            )
            merged = materialize_markdown_sidecars(
                run_dir,
                company="Acme",
                job_title="Role",
                job_id="na",
                job_date="20260708",
            )
            self.assertEqual(len(merged), 1)
            doc = Document(str(brief_docx))
            body = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("APPLY", body)

    def test_merge_brief_by_headings_when_markers_gone(self) -> None:
        """Pre-filled brief DOCX (no {{SECTION:}} markers) still merges via Heading 2 labels."""
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed")
        if not SCAFFOLD_BRIEF.is_file():
            self.skipTest("scaffold brief template missing")

        patch = """## OPPORTUNITY_INTELLIGENCE
Role: Staff Engineer at Calendly

## RECOMMENDATION
APPLY — strong platform leadership fit.
"""
        with tempfile.TemporaryDirectory() as tmp:
            dest = Path(tmp) / "brief.docx"
            copy_template(SCAFFOLD_BRIEF, dest)
            # Simulate legacy pre-merge: replace markers with stub text.
            doc = Document(str(dest))
            for para in doc.paragraphs:
                if "{{SECTION:" in para.text:
                    para.clear()
                    para.add_run("(KB stub — replaced)")
            doc.save(str(dest))

            count = merge_section_patch_into_docx(
                dest,
                parse_section_markdown(patch),
                bracket_ctx={
                    "company": "Calendly",
                    "job_title": "Staff Engineer",
                    "job_id": "na",
                    "job_date": "20260708",
                    "url": "https://example.com",
                },
                brief_heading_fallback=True,
            )
            self.assertGreater(count, 0)
            doc = Document(str(dest))
            body = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Calendly", body)
            self.assertIn("APPLY", body)
            # Patched sections replace stubs; unpatched sections may still show stub text.
            opp_idx = body.index("Opportunity Intelligence")
            rec_idx = body.index("Recommendation")
            self.assertIn("Staff Engineer at Calendly", body[opp_idx:rec_idx])


class CoachSidecarTests(unittest.TestCase):
    def test_coach_sidecar_status_prefers_patch(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            run_dir = Path(tmp)
            write_plain_docx(run_dir / "acme_role_na_20260708_resume.docx", paragraphs=["r"])
            (run_dir / "acme_role_na_20260708_resume_patch.md").write_text("## SUMMARY\nx", encoding="utf-8")

            status = coach_sidecar_status(run_dir)
            self.assertTrue(status["resume"])
            self.assertFalse(status["cover"])
            self.assertFalse(status["brief"])


if __name__ == "__main__":
    unittest.main()
